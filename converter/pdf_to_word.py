"""Módulo para conversão de PDF para Word.

Funcionalidades:
- Conversão de PDF para DOCX
- Preservação de formatação básica
- Suporte a PDFs com texto e escaneados (via OCR)
"""
from pathlib import Path
import logging
import tempfile

try:
    import docx
    from docx import Document
    from docx.shared import Inches
except ImportError:
    docx = None
    Document = None

try:
    from pdf2image.exceptions import PDFInfoNotInstalledError
except ImportError:
    PDFInfoNotInstalledError = None

from .ocr import extract_text_with_ocr, clean_text

logger = logging.getLogger(__name__)


def pdf_to_word(pdf_path: str, output_path: str = None, use_ocr: bool = False, clean_special_chars: bool = True) -> str:
    """Converte PDF para Word (DOCX).
    
    Args:
        pdf_path: Caminho para o arquivo PDF
        output_path: Caminho de saída (opcional, será gerado automaticamente se não fornecido)
        use_ocr: Se deve usar OCR para PDFs escaneados
        clean_special_chars: Se deve limpar caracteres especiais
    
    Returns:
        Caminho do arquivo DOCX criado
    
    Raises:
        RuntimeError: Se dependências não estão instaladas
    """
    if not docx:
        raise RuntimeError(
            "Dependência python-docx não encontrada. Instale: pip install python-docx"
        )
    
    pdf_path = Path(pdf_path)
    
    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    try:
        logger.info(f"Convertendo PDF para Word: {pdf_path} -> {output_path}")
        
        extraction_method = "Extração de texto"
        extraction_note = None

        # Extrair texto do PDF
        if use_ocr:
            try:
                text = extract_text_with_ocr(
                    str(pdf_path),
                    force_ocr=True,
                    clean_special_chars=clean_special_chars
                )
                extraction_method = "OCR"
            except Exception as ocr_error:
                logger.warning("OCR indisponível para %s: %s", pdf_path, ocr_error)
                extraction_note = str(ocr_error)
                if PDFInfoNotInstalledError and isinstance(ocr_error, PDFInfoNotInstalledError):
                    extraction_note = (
                        "Poppler não está instalado ou não está no PATH. "
                        "Continuando com extração de texto padrão."
                    )
                extraction_method = "OCR (fallback para texto)"
                from .reader import read_pdf  # import local to evitar ciclo
                text = read_pdf(pdf_path)
                if clean_special_chars:
                    text = clean_text(text)
        else:
            # Tentar extração normal primeiro, depois OCR se necessário
            text = extract_text_with_ocr(
                str(pdf_path),
                force_ocr=False,
                clean_special_chars=clean_special_chars
            )
        
        # Criar documento Word
        doc = Document()
        
        # Adicionar título
        title = doc.add_heading(f'Documento convertido: {pdf_path.name}', 0)
        
        # Adicionar informações do arquivo
        info_para = doc.add_paragraph()
        info_para.add_run('Arquivo original: ').bold = True
        info_para.add_run(str(pdf_path))
        
        info_para = doc.add_paragraph()
        info_para.add_run('Método de extração: ').bold = True
        info_para.add_run(extraction_method)
        if extraction_note:
            note_para = doc.add_paragraph()
            note_para.add_run('Observação: ').bold = True
            note_para.add_run(extraction_note)
        
        # Adicionar separador
        doc.add_paragraph('─' * 50)
        
        # Adicionar conteúdo
        if text.strip():
            # Dividir texto em parágrafos
            paragraphs = text.split('\n\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    # Verificar se é um título (linha curta, sem pontuação final)
                    if len(para_text) < 100 and not para_text.endswith(('.', '!', '?', ':')):
                        doc.add_heading(para_text, level=1)
                    else:
                        doc.add_paragraph(para_text)
        else:
            doc.add_paragraph("Nenhum texto foi extraído do documento.")
        
        # Salvar documento
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        logger.info(f"Conversão concluída: {output_path}")
        return str(output_path)
        
    except Exception as e:
        logger.error(f"Erro durante conversão PDF para Word: {e}")
        raise


def batch_pdf_to_word(pdf_paths: list, output_dir: str, use_ocr: bool = False, clean_special_chars: bool = True) -> list:
    """Converte múltiplos PDFs para Word.
    
    Args:
        pdf_paths: Lista de caminhos para arquivos PDF
        output_dir: Diretório de saída
        use_ocr: Se deve usar OCR para PDFs escaneados
        clean_special_chars: Se deve limpar caracteres especiais
    
    Returns:
        Lista de caminhos dos arquivos DOCX criados
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    converted_files = []
    
    for pdf_path in pdf_paths:
        try:
            pdf_path = Path(pdf_path)
            output_path = output_dir / f"{pdf_path.stem}.docx"
            
            result_path = pdf_to_word(
                str(pdf_path),
                str(output_path),
                use_ocr=use_ocr,
                clean_special_chars=clean_special_chars
            )
            
            converted_files.append(result_path)
            
        except Exception as e:
            logger.error(f"Erro convertendo {pdf_path}: {e}")
            continue
    
    return converted_files